# app.py
import os
import re
import textwrap
from io import BytesIO

import streamlit as st
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------- App Config ----------
st.set_page_config(page_title="Sinapis AI Coach – BMC Review", page_icon="🧭", layout="wide")
st.title("Sinapis AI Coach – Ascent BMC Review")
st.markdown("Upload the **Word submission** and click **Run Review**. You’ll get a downloadable Word report.")

# ---------- OpenAI (lazy client + firm timeout) ----------
def get_client():
    api_key = st.secrets.get("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY")
    org_id  = st.secrets.get("OPENAI_ORG")     or os.getenv("OPENAI_ORG") or os.getenv("OPENAI_ORG_ID")
    proj_id = st.secrets.get("OPENAI_PROJECT") or os.getenv("OPENAI_PROJECT")
    if not api_key:
        st.error("Missing OPENAI_API_KEY in Streamlit Secrets.")
        st.stop()
    os.environ["OPENAI_API_KEY"] = api_key
    if org_id:  os.environ["OPENAI_ORG_ID"] = org_id
    if proj_id: os.environ["OPENAI_PROJECT"] = proj_id
    return OpenAI(timeout=60.0)

MODEL_NAME = "gpt-4o" if (os.getenv("USE_GPT4O") == "1" or st.secrets.get("USE_GPT4O") == "1") else "gpt-4o-mini"
MAX_TOKENS = 7000 if MODEL_NAME == "gpt-4o" else 5000

# ---------- Prompts ----------
SINAPIS_COACH_SYS = textwrap.dedent("""
You are **Sinapis AI Coach**, reviewing founder submissions using the Sinapis Ascent Business Model Canvas.
Audience: post-revenue SMEs in frontier markets (Kenya first).
Method: Osterwalder BMC + Lean Canvas emphasis + Sinapis Kingdom Impact lens.
Tone: encouraging, direct critique; diagnostic + light prescriptive.
NEVER invent content for missing blocks; mark “Missing/Needs input.”
""").strip()

MARKDOWN_INSTRUCTION = (
    "Return the assessment as Markdown. "
    "Use '##' for major sections (1) Problem … 14) Final Assessment). "
    "Use '###' for subheadings (Strengths, Weaknesses, Probing Questions, Suggested Explorations; "
    "and under Final Assessment: Quick Wins, Deeper Strategic Questions, Overall Cohesiveness). "
    "Use bullet lists for items under each subheading. Do not use bold for subheadings."
)

STRICT_NO_INVENTION = (
    "CRITICAL: Base the assessment ONLY on the JSON fields provided. "
    "For ANY block whose input is empty/whitespace, mark the block as Missing/Needs input. "
    "Under each of its subheadings, output a single bullet: 'Missing/Needs input.' "
    "Do NOT infer or fabricate. This review is stateless and only for this submission."
)

# --- Depth defaults ---
DEPTH_MIN_COUNTS = {"Strengths": 4, "Weaknesses": 6, "Probing Questions": 10, "Suggested Explorations": 8}

KENYA_LENS = """
Contextualize critiques for Kenya/East Africa where relevant:
- Mobile money (e.g., M-Pesa), agent networks, last-mile logistics, power/connectivity reliability.
- Seasonality (agri, school terms), cash cycle & working capital constraints, FX risk, import duties/customs.
- County-level regulation & permits, KEBS/product standards, data privacy basics.
"""

DEPTH_INSTRUCTION = f"""
Depth Mode (default) — Be rigorous and specific while staying diagnostic (no company-specific step-by-step).
For each major section:
- Begin with a compact score 'Score: X/5' + 1-line reason.
- **Strengths**: ≥ {DEPTH_MIN_COUNTS['Strengths']} bullets spanning: Market, Customer, Competition, Ops, Finance, Impact, Team/Governance.
- **Weaknesses**: ≥ {DEPTH_MIN_COUNTS['Weaknesses']} bullets; call out evidence gaps/assumptions.
- **Probing Questions**: ≥ {DEPTH_MIN_COUNTS['Probing Questions']} bullets; cover Market, Customer, Competition, Ops, Finance/Unit economics, Impact/ESG, Legal/Regulatory, Distribution.
- **Suggested Explorations**: ≥ {DEPTH_MIN_COUNTS['Suggested Explorations']} bullets; use experiment categories (smoke test, concierge/pilot, pricing, channel trial, churn interview, service blueprint, instrumentation).
If a block is empty, keep only 'Missing/Needs input.' as required by the strict rule.
"""

CONSISTENCY_MATRIX = """
Include cross-check bullets where relevant:
- Value↔Segment fit, Segment↔Channels reach/cost, Costs↔Revenue seasonality/cash cycle, Activities↔Resources & Partners.
If numbers are missing for unit economics, state the exact numbers required (price, gross margin %, CAC, churn %, payback).
"""

# ✅ NEW: response template constant (replaces the broken inline string)
SINAPIS_RESPONSE_TEMPLATE = textwrap.dedent("""
Use exactly these headings and order:

1) Problem
   Strengths
   Weaknesses
   Probing Questions
   Suggested Explorations
2) Value Proposition
   Strengths
   Weaknesses
   Probing Questions
   Suggested Explorations
3) Unfair Advantage
   Strengths
   Weaknesses
   Probing Questions
   Suggested Explorations
4) Customer Segments
   Strengths
   Weaknesses
   Probing Questions
   Suggested Explorations
5) Channels
   Strengths
   Weaknesses
   Probing Questions
   Suggested Explorations
6) Customer Relationships
   Strengths
   Weaknesses
   Probing Questions
   Suggested Explorations
7) Key Activities
   Strengths
   Weaknesses
   Probing Questions
   Suggested Explorations
8) Key Resources
   Strengths
   Weaknesses
   Probing Questions
   Suggested Explorations
9) Key Partners
   Strengths
   Weaknesses
   Probing Questions
   Suggested Explorations
10) Revenue Streams
   Strengths
   Weaknesses
   Probing Questions
   Suggested Explorations
11) Cost Structure
   Strengths
   Weaknesses
   Probing Questions
   Suggested Explorations
12) Kingdom Impact
   Strengths
   Weaknesses
   Probing Questions
   Suggested Explorations

13) Cross-Block Observations
   Inconsistencies
   Opportunities to Strengthen

14) Final Assessment
   Quick Wins
   Deeper Strategic Questions
   Overall Cohesiveness

Footer: Advisory—Not Legal/Financial Advice.
""").strip()

# --- Optional: load rubric/workbook snippets from guides/ if present ---
def read_guide_if_exists(rel_path: str, max_chars: int = 10000) -> str:
    base_dir = os.path.dirname(__file__)
    path = os.path.join(base_dir, rel_path)
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                return f.read()[:max_chars]
        except Exception:
            return ""
    return ""

RUBRIC_GUIDE = read_guide_if_exists(os.path.join("guides", "sinapis_rubric.md"))
WORKBOOK_GUIDE = read_guide_if_exists(os.path.join("guides", "sinapis_workbook.md"))

# (Optional debug; remove later)
if RUBRIC_GUIDE:
    st.caption(f"Loaded rubric guide ({len(RUBRIC_GUIDE)} chars)")
else:
    st.caption("Rubric guide not found")
if WORKBOOK_GUIDE:
    st.caption(f"Loaded workbook guide ({len(WORKBOOK_GUIDE)} chars)")
else:
    st.caption("Workbook guide not found")

# ---------- Parsing helpers (unchanged) ----------
def normalize_label(s: str) -> str:
    if not s: return ""
    s = s.strip()
    s = re.sub(r"^\s*\d+\s*[\.\)\-:]?\s*", "", s)
    s = s.rstrip(":").strip().lower()
    s = re.sub(r"\s+", " ", s)
    return s

FIELD_ALIASES = {
    "business_name": ["business name"],
    "brief_description": ["brief description of business", "brief description"],
    "problem": ["problem","1) problem","1. problem"],
    "value_proposition": ["value proposition","2) value proposition","2. value proposition"],
    "unfair_advantage": ["unfair advantage","3) unfair advantage","3. unfair advantage"],
    "customer_segments": ["customer segments","4) customer segments","4. customer segments"],
    "channels": ["channels","5) channels","5. channels"],
    "customer_relationships": ["customer relationships","6) customer relationships","6. customer relationships"],
    "key_activities": ["key activities","7) key activities","7. key activities"],
    "key_resources": ["key resources","8) key resources","8. key resources"],
    "key_partners": ["key partners","9) key partners","9. key partners"],
    "revenue_streams": ["revenue streams","10) revenue streams","10. revenue streams"],
    "cost_structure": ["cost structure","11) cost structure","11. cost structure"],
    "kingdom_impact": ["kingdom impact","12) kingdom impact","12. kingdom impact"],
}
ALIAS_TO_KEY = {normalize_label(a): k for k, arr in FIELD_ALIASES.items() for a in arr}

BASE_PHRASE = {
    "business_name": "business name",
    "brief_description": "brief description",
    "problem": "problem",
    "value_proposition": "value proposition",
    "unfair_advantage": "unfair advantage",
    "customer_segments": "customer segments",
    "channels": "channels",
    "customer_relationships": "customer relationships",
    "key_activities": "key activities",
    "key_resources": "key resources",
    "key_partners": "key partners",
    "revenue_streams": "revenue streams",
    "cost_structure": "cost structure",
    "kingdom_impact": "kingdom impact",
}

def guess_key_from_label_cell(left_text: str):
    if not left_text:
        return None
    lines = [normalize_label(x) for x in left_text.splitlines()]
    lines = [x for x in lines if x]
    for ln in lines:
        if ln in ALIAS_TO_KEY:
            return ALIAS_TO_KEY[ln]
    for ln in lines:
        for key, phrase in BASE_PHRASE.items():
            if phrase in ln:
                return key
    return None

HINT_SNIPPETS = [
    "provide a brief description",
    "what customer problem",
    "what are you offering",
    "what is your uniqueness",
    "which customer groups",
    "through what means do you reach",
    "what type of relationship",
    "what tasks are vital",
    "what assets are essential",
    "which external organizations",
    "how does your business earn revenue",
    "what are the defining characteristics of your cost structure",
    "where and how are you intentionally looking to make impact",
]

def clean_value(text: str) -> str:
    lines = [ln.strip() for ln in (text or "").splitlines()]
    out = []
    for ln in lines:
        if not ln: continue
        if any(h in ln.lower() for h in HINT_SNIPPETS): continue
        out.append(ln)
    return "\n".join(out).strip()

def parse_docx_to_payload(doc_bytes: bytes) -> dict:
    doc = Document(BytesIO(doc_bytes))
    buf = {k: "" for k in FIELD_ALIASES.keys()}
    saw_nonempty = False
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2: continue
            key = guess_key_from_label_cell(row.cells[0].text)
            val = clean_value(row.cells[1].text)
            if key and val:
                buf[key] = val
                saw_nonempty = True
    if saw_nonempty:
        return buf
    current_key = None
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t: continue
        norm = normalize_label(t)
        if norm in ALIAS_TO_KEY:
            current_key = ALIAS_TO_KEY[norm]; continue
        if current_key:
            if norm in ALIAS_TO_KEY:
                current_key = ALIAS_TO_KEY[norm]; continue
            val = clean_value(t)
            if val:
                buf[current_key] = (buf[current_key] + "\n" + val).strip()
    return buf

# ---------- Markdown normalization & enforcement ----------
MAJOR_SECTIONS = [
    "1) Problem","2) Value Proposition","3) Unfair Advantage","4) Customer Segments",
    "5) Channels","6) Customer Relationships","7) Key Activities","8) Key Resources",
    "9) Key Partners","10) Revenue Streams","11) Cost Structure","12) Kingdom Impact",
    "13) Cross-Block Observations","14) Final Assessment"
]
SUBS_STANDARD = ["Strengths","Weaknesses","Probing Questions","Suggested Explorations"]
SUBS_13 = ["Inconsistencies","Opportunities to Strengthen"]
SUBS_14 = ["Quick Wins","Deeper Strategic Questions","Overall Cohesiveness"]
ADVISORY_TEXTS = {"Advisory—Not Legal/Financial Advice.","Footer: Advisory—Not Legal/Financial Advice."}

def list_empty_blocks(payload: dict):
    mapping = {
        "1) Problem":"problem","2) Value Proposition":"value_proposition","3) Unfair Advantage":"unfair_advantage",
        "4) Customer Segments":"customer_segments","5) Channels":"channels","6) Customer Relationships":"customer_relationships",
        "7) Key Activities":"key_activities","8) Key Resources":"key_resources","9) Key Partners":"key_partners",
        "10) Revenue Streams":"revenue_streams","11) Cost Structure":"cost_structure","12) Kingdom Impact":"kingdom_impact",
    }
    return [title for title, k in mapping.items() if not (payload.get(k) or "").strip()]

def normalize_markdown(md_text: str) -> str:
    out = []
    for raw in md_text.splitlines():
        line = raw.strip()
        if not line: out.append(""); continue
        if line in ADVISORY_TEXTS: continue
        if line in MAJOR_SECTIONS: out.append("## " + line); continue
        if line in (SUBS_STANDARD + SUBS_13 + SUBS_14): out.append("### " + line); continue
        out.append(line)
    return "\n".join(out).strip()

def enforce_missing_for_empty_blocks(norm_md: str, empty_blocks: list[str]) -> str:
    lines = norm_md.splitlines()
    idx = [(i, lines[i][3:].strip()) for i in range(len(lines)) if lines[i].startswith("## ")]
    idx.append((len(lines), None))
    def subs_for(t): return SUBS_13 if t=="13) Cross-Block Observations" else (SUBS_14 if t=="14) Final Assessment" else SUBS_STANDARD)
    out, i = [], 0
    while i < len(lines):
        if lines[i].startswith("## "):
            title = lines[i][3:].strip()
            next_pos = next((k for k,(pos,_) in enumerate(idx) if pos==i), None)
            end = idx[next_pos+1][0] if next_pos is not None else len(lines)
            if title in empty_blocks:
                out.append(lines[i])
                for s in subs_for(title):
                    out.append(f"### {s}")
                    out.append("• Missing/Needs input.")
                i = end; continue
        out.append(lines[i]); i += 1
    return "\n".join(out).strip()

# ---------- DOCX builder ----------
def build_docx_from_markdown(md_text: str, founder_payload: dict) -> bytes:
    doc = Document()
    base_dir = os.path.dirname(__file__)
    logo_path = os.path.join(base_dir, "assets", "logo.png")
    if os.path.exists(logo_path):
        try:
            p = doc.add_paragraph(); r = p.add_run()
            r.add_picture(logo_path, width=Inches(1.5)); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception: pass
    title = doc.add_heading(f"Sinapis AI Coach – BMC Review of {founder_payload.get('business_name') or '(Unnamed Business)'}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bd = founder_payload.get("brief_description") or "—"
    meta = doc.add_paragraph(); r1 = meta.add_run("Description: "); r1.bold = True; meta.add_run(bd)
    normal = doc.styles["Normal"].font; normal.name = "Calibri"; normal.size = Pt(11)
    h1 = doc.styles["Heading 1"].font; h1.name = "Calibri"; h1.size = Pt(14); h1.bold = True; h1.color.rgb = RGBColor(31,78,121)
    h2 = doc.styles["Heading 2"].font; h2.name = "Calibri"; h2.size = Pt(12); h2.bold = True; h2.color.rgb = RGBColor(0,0,0)
    for raw in md_text.splitlines():
        line = raw.strip()
        if line == "": doc.add_paragraph(""); continue
        if line.startswith("## "):  doc.add_heading(line[3:].strip(), level=1); continue
        if line.startswith("### "): doc.add_heading(line[4:].strip(), level=2); continue
        if line.startswith(("• ","- ")): doc.add_paragraph(line[2:].strip(), style="List Bullet"); continue
        doc.add_paragraph(line)
    doc.add_paragraph().add_run("Advisory—Not Legal/Financial Advice.").italic = True
    buf = BytesIO(); doc.save(buf); buf.seek(0); return buf.getvalue()

def render_download_only(markdown_text: str, founder_payload: dict):
    st.success("Your AI review is ready. Click below to download the Word report.")
    st.download_button(
        label="⬇️ Download your review (Word .docx)",
        data=build_docx_from_markdown(markdown_text, founder_payload),
        file_name="Sinapis_AI_Coach_Assessment.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

# ---------- UI ----------
uploaded = st.file_uploader("Upload founder submission (.docx)", type=["docx"])
submitted = st.button("Run Review", use_container_width=True, disabled=uploaded is None)

# ---------- Run Review ----------
if submitted and uploaded:
    with st.spinner("Parsing your submission…"):
        try:
            payload = parse_docx_to_payload(uploaded.read())
        except Exception as e:
            st.error(f"Failed to read .docx: {e}"); st.stop()

        filled_blocks = sum(1 for k in [
            "problem","value_proposition","unfair_advantage","customer_segments",
            "channels","customer_relationships","key_activities","key_resources",
            "key_partners","revenue_streams","cost_structure","kingdom_impact"
        ] if (payload.get(k) or "").strip())
        st.info(f"Parsed {filled_blocks}/12 canvas blocks from the upload.")
        if filled_blocks == 0:
            st.error("No canvas content detected. Ensure your file has a two-column table with section names in the left column and your input in the right column, or use our template.")
            st.stop()

    with st.spinner("Contacting OpenAI…"):
        empty_blocks = list_empty_blocks(payload)
        user_message = (
            "Founder Input (normalized JSON):\n"
            + str(payload)
            + "\n\nEMPTY_BLOCKS: " + str(empty_blocks)
            + "\n\nUse the response template exactly."
        )

        # Build messages (deep critique by default + optional guides)
        messages = [
            {"role": "system", "content": SINAPIS_COACH_SYS},
            {"role": "system", "content": MARKDOWN_INSTRUCTION},
            {"role": "system", "content": STRICT_NO_INVENTION},
            {"role": "system", "content": DEPTH_INSTRUCTION},
            {"role": "system", "content": KENYA_LENS},
            {"role": "system", "content": CONSISTENCY_MATRIX},
            {"role": "system", "content": "Response Template:\n" + SINAPIS_RESPONSE_TEMPLATE},
        ]
        if RUBRIC_GUIDE:
            messages.insert(0, {"role": "system", "content": "Sinapis Internal Rubric (excerpt):\n" + RUBRIC_GUIDE})
        if WORKBOOK_GUIDE:
            messages.insert(1, {"role": "system", "content": "Sinapis Workbook Notes (excerpt):\n" + WORKBOOK_GUIDE})
        messages.append({"role": "user", "content": user_message})

        try:
            client = get_client()
            resp = client.chat.completions.create(
                model=MODEL_NAME,
                temperature=0.0,
                max_tokens=MAX_TOKENS,
                messages=messages,
            )
            raw_md = resp.choices[0].message.content
        except Exception as e:
            st.error(f"OpenAI request failed (did not complete): {e}"); st.stop()

    with st.spinner("Finalizing your report…"):
        norm_md = normalize_markdown(raw_md)
        final_md = enforce_missing_for_empty_blocks(norm_md, empty_blocks)
        render_download_only(final_md, payload)

# ---------- Footer ----------
st.caption("Advisory—Not Legal/Financial Advice.")
